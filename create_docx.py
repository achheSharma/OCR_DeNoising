"""
This script contains the main logic to convert document
from parsed xml file which is created from poppler library for text extraction.

The complete conversion process involves following steps in this script:
    
    * parsing xml document in the form of xml tree from ElementTree.
    
    * This parsed document is read in DFS manner, OCR denoising techniques
      are applied onto it and stored in list structure containing nested dictionaries
      to store data in key-value pair for fast accessing of data.
      Note: For constructing this resultant data structure some minor changes are
            done to xml while executing DFS function.
    
    * This resultant data structure is then used to construct word document.
      
Techniques used for OCR Denoising:
    
    * SpellChecker for detection of wrong words and their conversion to correct ones 
      with most probable substitute word trial dictionary specified below.
    
    * A list of most common character conversion errors in words as well as digits.
    
    * Maintaining a list of noun words to be omitted from check and correction matching
      purposes.

Current denoising techniques can be made better with adding more such common errors, trials and
list of nouns that won't be checked against spelling checker.

This script contains the following functions:
    
    * hasNumbers - check if the passed string containing any digits or not.
    * correctNums - for correcting most common digits related issues with poppler
                    library keeping current use case in consideration.
    * checkErrors - Recursively substitute characters in a word for checking 
                    against most plausable correct words with calculation of 
                    Jaccard distance logic.
    * spellCheck - It extract given word by removing starting and terminating
                   punctuations from a given word parsed from parseText.
    * parseText - It parses passed text onto other methods for OCR denoising process.
    * style - paragraph styling in docx specifying styling attributes like bold, italics,
              underline, font family and size extraction.
    * DFS - It iterates over complete xml tree constructed in a depth first manner for
            creating storing data structure for creation of word document.
    * create_docx - Main logic for converting given list data structure into a docx.

"""

import os
import re
import string
import xml.etree.ElementTree as ET
# for this import of module install python docx library
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
# for this import of module install pyspellchecker library
from spellchecker import SpellChecker


# OCR denoising techniques
erroneous = {'AH': 'All', 'S': '$', 's': '$', '<E':'(£'}
trial = {'l':'t', 'H':'ll' }
numericalErrors = {'O':'0', 'o':'0', 'I':'1', 'l':'1'}
currency = ["$", "£"]
# Pyspellchecker for suggesting alternate correct words
# with minimum Jaccard distance of alternate suggested or current word.
spell  = SpellChecker()
# List of nouns to be omitteed from SpellChecker module and
# assumed to be correct and compared against only 'words.txt' file.
spell.word_frequency.load_text_file('./words.txt')

def hasNumbers(inputString):
    """
    This function checks if inputted string contains digits or not.
    
    Parameters
    ----------
    inputString : str
        input string to be checked.
        
    Returns
    -------
    boolean
        either True or False depending upon digits passed.
    """
    return any(char.isdigit() for char in inputString)

def correctNums(inputString):
    """
    This function checks for frequent errors in digits
    and given suitable changes accordingly.
    
    Parameters
    ----------
    inputString : str
        input string to be checked.

    Returns
    -------
    str
        returns correct word
    """
    outputString = ""
    for i in range(len(inputString)):
        if inputString[i] in numericalErrors.keys():
            outputString = outputString + numericalErrors[inputString[i]]
        else:
            outputString = outputString + inputString[i]
    return outputString

def checkErrors(word, n):
    """
    This is a recursive function checks if the current word is meaningful
    otherwise it substitute most probable misread individual characters with
    suitable correction characters in recursive manner. 
    
    Suitable correct word is returned with minimum Jaccard distance from
    existing vocabulary of correct words.
    
    Parameters
    ----------
    word : str
        input string for given word to be checked.
    n : int
        the character that will be changed in that particular call.
        
    Returns
    -------
    boolean
        either True or False depending upon new word has error or not.
    str
        suitable corrected or original word.            
    """
    if word in erroneous.keys():
        return erroneous[word]
    correction = spell.correction(word.lower())
    if correction.lower() != word.lower():
        w = word
        for i in range(n, len(word)):
            if word[i] in trial.keys():
                word = word[:i] + trial[word[i]] + word[i+1:]
                flag, newWord = checkErrors(word, i+1)
                if flag:
                    return True, newWord
                else:
                    word = w
    else:
        return True, word
    return False, word

def spellCheck(word):
    """
    It extract given word by removing starting and terminating
    punctuations from a given word parsed from parseText.
    
    Parameters
    ----------
    word : str
        input string for given word to be checked.
    
    Returns
    -------
    str
        correct word with complete punctuations returned.
    """
    tempWord = word
    startPunc = ""
    for i in range(len(word)):
        if word[i] in string.punctuation:
            startPunc = startPunc+word[i]
            tempWord = tempWord[1:]
        else:
            break

    endPunc = ""
    for i in range(len(word)-1, -1, -1):
        if word[i] in string.punctuation:
            endPunc = word[i]+endPunc
            tempWord = tempWord[:-1]
        else:
            break

    title = tempWord.istitle()
    upper = tempWord.isupper()
    word = tempWord
    
    try:
        flag, word = checkErrors(word, 0)
    except:
        pass

    if title:
        word = word.title()
    elif upper:
        word = word.upper()
    return startPunc + word + endPunc

def parseText(text):
    """
    It parses passed text onto other methods for OCR denoising process.
    first it splits the text passed and checks its spelling with spellCheck
    and checkErrors functions.

    Parameters
    ----------
    text : str
        input string containing multiple words.
    
    Returns
    -------
    text
        correct words with complete punctuations returned.
    """
    text = text.split()
    for i, t in enumerate(text):
        if t in currency:
            continue
        elif t in erroneous.keys():
            text[i] = erroneous[t]
        elif hasNumbers(t):
            text[i] = correctNums(t)
        else:
            text[i] = spellCheck(t)
    text = ' '.join(text)
    return text

# Parameters for storing docx related information from given xml.
fonts = {}
prevTop = None
fontStyle = [False, False, False] # Bold, Italics, Underline
pageHeight = None
pageWidth = None
leftMargin = None

def DFS(parent, page_data, fontStyle):
    """
    It iterates over complete xml tree constructed in a depth first manner for
    creating storing data structure for creation of word document.
    
    Parameters
    ----------
    parent : object of tree
        parent of given node.
    page_data : dictionary
        key-value pair information about attributes and data as key
        of particular page.
    fontStyle : list
        different styles of font either bold, italics or underline.
    
    
    Returns
    -------
    str
        text returned as string of every object.
    list
        list of size three of different font styles.
    """
    
    text = None
    for elem in parent:
        fontStyle = [False, False, False]
        if elem.tag == 'fontspec':
            fonts[elem.attrib['id']] = elem.attrib
            continue
        text, fontStyle = DFS(elem, page_data,fontStyle)
        if not text:
            text = elem.text
        if elem.tag == 'b':
            fontStyle[0] = True
            return text, fontStyle
        elif elem.tag == 'i':
            fontStyle[1] = True
            return text, fontStyle
        elif elem.tag == 'u':
            fontStyle[2] = True
            return text, fontStyle

        if elem.tag == 'text':
            if not text or text == ' ':
                continue
            attrib = elem.attrib
            attrib['b'] = fontStyle[0]
            attrib['i'] = fontStyle[1]
            attrib['u'] = fontStyle[2]
            # leftMargin = min(leftMargin, int(attrib['left']))
            if elem.tag in page_data:
                text = parseText(text)
                page_data[elem.tag].append([text, attrib])
            else:
                page_data[elem.tag] = [[text, attrib]]
    return text, fontStyle

def style(run, text, attrib, prevLeft = None):
    """
    paragraph styling in docx specifying styling attributes like bold, italics,
    underline, font family and size extraction.
    
    Parameters
    ----------
    run : function
        function that formats data on the fly as data is passed.
    text : str
        text passed as input for being inserted along with given style.
    attrib : dictionary
        storing all the styles of the text that is passed.
    prevLeft : int
        numeric value for tab prediction and calculation in current running paragraph.
    """
    
    font = run.font
    tab = ""
    if prevLeft:
        diff = abs(int(attrib['left']) - prevLeft)
        for i in range(int(diff/50)):
            tab = tab + "\t"
    run.text = tab + text + " "
    font.name = fonts[attrib['font']]['family']
    font.size = Pt(int(fonts[attrib['font']]['size']))
    if 'b' in attrib.keys():
        font.bold = attrib['b']
    if 'i' in attrib.keys():
        font.italic = attrib['i']
    if 'u' in attrib.keys():
        font.underline = attrib['u']



def create_docx(path, fileName):
    """
    Main logic for converting given list data structure into a docx.
    
    For conversion page is selected from given created data structure.
    Line by line text is inserted in sorted manner in top-down and
    left-right manner. Correction functions, tab insertion logic are
    also called from this function.
    
    Parameters
    ----------
    path : str
        xml file's path that is to be passed for conversion.
    fileName : str
        name of file to be passed.
    """

    xml = path + fileName
    fileName = fileName[:-8]
    tree  = ET.parse(xml)
    root = tree.getroot()

    pages = []
    for page in root:
        page_data = {}
        fontStyle = [False, False, False]
        pageHeight = int(page.attrib['height'])
        pageWidth = int(page.attrib['width'])
        leftMargin = pageWidth
        DFS(page, page_data, fontStyle)
        pages.append(page_data)

    document = Document()

    count = 0
    for page in pages:
        top = None
        for key, value in page.items():
            if key == 'text':
                value.sort(key = lambda x: int(x[1]['top']))
                for i in range(len(value)-1):
                    if abs(int(value[i][1]['top']) - int(value[i+1][1]['top'])) <= 5:
                        value[i+1][1]['top'] = value[i][1]['top']
                value.sort(key = lambda x: (int(x[1]['top']), int(x[1]['left'])))
                prevLeft = None
                prevTop = None
                align = None
                for item in value:
                    text = item[0]
                    attrib = item[1]
                    if top != attrib['top']:
                        center = False
                        top = attrib['top']
                        left = int(attrib['left'])
                        para = document.add_paragraph()
                        para_format = para.paragraph_format
                        if prevTop:
                            if abs(int(top) - prevTop) < 25:
                                para_format.space_after = Pt(0)
                        run = para.add_run()
                        style(run, text, attrib)
                    else:
                        run = para.add_run()
                        style(run, text, attrib, prevLeft)
                    if left > 200 and abs(pageWidth - (left + int(attrib['width'])) - int(attrib['left'])) < 50:
                            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            align = "center"
                    if left > pageWidth/2 and int(attrib['left']) + int(attrib['width']) > pageWidth - 300:
                        para_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        align = "right"

                    prevLeft = int(attrib['left']) + int(attrib['width'])
                    prevTop = int(attrib['top']) + int(attrib['height'])
        count = count+1
        if count < len(pages):
            document.add_page_break()
    try:
        os.makedirs("Docx/" + fileName)
    except:
        pass
    document.save("Docx/" + fileName + "/" + fileName + ".docx")
