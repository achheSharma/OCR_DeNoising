import os
import copy
import re
import string

import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from pprintpp import pprint as pp  # pretty print karane ke liye dictionary ko

from spellchecker import SpellChecker

erroneous = {'AH': 'All', 'S': '$', 's': '$'}
trial = {'l':'t', 'H':'ll'}
currency = ["$"]

spell  = SpellChecker()
spell.word_frequency.load_text_file('./words.txt')

def hasNumbers(inputString):
    return any(char.isdigit() for char in inputString)

def checkCurr(inputString):
    inputString = inputString.split(',')
    for i in inputString:
        if not i.isdigit():
            return False
    return True

def checkDate(inputString):
    return True

def checkErrors(word, n):
    if word in erroneous.keys():
        return erroneous[word]
    correction = spell.correction(word.lower())
    if correction.lower() != word.lower():
        # print(spell.candidates(correction))
        w = word
        for i in range(n, len(word)):
            if word[i] in trial.keys():
                word = word[:i] + trial[word[i]] + word[i+1:]
                flag, newWord = checkErrors(word, i+1)
                if flag:
                    return True, newWord
                else:
                    word = w
    else:
        return True, word
    return False, word

def spellCheck(word):
    tempWord = word
    startPunc = ""
    for i in range(len(word)):
        if word[i] in string.punctuation:
            startPunc = startPunc+word[i]
            tempWord = tempWord[1:]
        else:
            break

    endPunc = ""
    for i in range(len(word)-1, -1, -1):
        if word[i] in string.punctuation:
            endPunc = word[i]+endPunc
            tempWord = tempWord[:-1]
        else:
            break

    title = tempWord.istitle()
    upper = tempWord.isupper()
    word = tempWord
    
    flag, word = checkErrors(word, 0)

    if title:
        word = word.title()
    elif upper:
        word = word.upper()
    return startPunc + word + endPunc

def parseText(text):
    text = text.split()
    # print(text)
    for i, t in enumerate(text):
        # print(t)
        if t in currency:
            continue
        elif t in erroneous.keys():
            text[i] = erroneous[t]
        elif hasNumbers(t):
            if ',' in t:
                if checkCurr(t):
                    pass
                    # print("Currency\n")
            elif '/' in t:
                if checkDate(t):
                    pass
                    # print("Date\n")
        else:
            text[i] = spellCheck(t)
    text = ' '.join(text)
    return text

fonts = {}
prevTop = None
fontStyle = [False, False, False] # Bold, Italics, Underline
pageHeight = None
pageWidth = None
leftMargin = None
def DFS(parent, page_data, fontStyle):
    # global leftMargin
    for elem in parent:
        if elem.tag == 'fontspec':
            fonts[elem.attrib['id']] = elem.attrib
            continue

        # if elem.tag == 'b':
        #     return elem.tag, elem.text
        text, fontStyle = DFS(elem, page_data, copy.deepcopy(fontStyle))
        if not text:
            text = elem.text
        if elem.tag == 'b':
            fontStyle[0] = True
            return text, fontStyle
        elif elem.tag == 'i':
            fontStyle[1] = True
            return text, fontStyle
        elif elem.tag == 'u':
            fontStyle[2] = True
            return text, fontStyle

        if elem.tag == 'text':
            if not text or text == ' ':
                continue
            attrib = elem.attrib
            attrib['b'] = fontStyle[0]
            attrib['i'] = fontStyle[1]
            attrib['u'] = fontStyle[2]
            # leftMargin = min(leftMargin, int(attrib['left']))
            if elem.tag in page_data:
                text = parseText(text)
                page_data[elem.tag].append([text, attrib])
            else:
                page_data[elem.tag] = [[text, attrib]]
    return None, fontStyle

def style(run, text, attrib, prevLeft = None):
    # print(text)
    font = run.font
    tab = ""
    if prevLeft:
        diff = abs(int(attrib['left']) - prevLeft)
        # print("\tTAB: ", diff)
        for i in range(int(diff/50)):
            tab = tab + "\t"
    run.text = tab + text + " "
    font.name = fonts[attrib['font']]['family']
    font.size = Pt(int(fonts[attrib['font']]['size']))
    if 'b' in attrib.keys():
        font.bold = attrib['b']
    if 'i' in attrib.keys():
        font.italic = attrib['i']
    if 'u' in attrib.keys():
        font.underline = attrib['u']

# pp(pages)

def createDocx(path, fileName):
    # global leftMargin
    print("Creating Docx")
    xml = path + fileName
    fileName = fileName[:-8]
    tree  = ET.parse(xml)
    root = tree.getroot()

    pages = []
    for page in root:
        page_data = {}
        fontStyle = [False, False, False]
        pageHeight = int(page.attrib['height'])
        pageWidth = int(page.attrib['width'])
        leftMargin = pageWidth
        DFS(page, page_data, fontStyle)
        pages.append(page_data)
    # pp(pages)

    # check image addition code ???
    # https://stackoverflow.com/questions/32932230/add-an-image-in-a-specific-position-in-the-document-docx-with-python
    # loop structure inconsistent

    document = Document()

    for page in pages:
        top = None
        for key, value in page.items():
            if key == 'text':
                value.sort(key = lambda x: int(x[1]['top']))
                for i in range(len(value)-1):
                    if abs(int(value[i][1]['top']) - int(value[i+1][1]['top'])) <= 5:
                        value[i+1][1]['top'] = value[i][1]['top']
                value.sort(key = lambda x: (int(x[1]['top']), int(x[1]['left'])))
                # pp(value)
                prevLeft = None
                prevTop = None
                center = False
                for item in value:
                    text = item[0]
                    attrib = item[1]
                    if top != attrib['top']:
                        center = False
                        top = attrib['top']
                        left = int(attrib['left'])
                        para = document.add_paragraph()
                        para_format = para.paragraph_format
                        if prevTop:
                            if abs(int(top) - prevTop) < 25:
                                para_format.space_after = Pt(0)
                        run = para.add_run()
                        style(run, text, attrib)
                    else:
                        run = para.add_run()
                        style(run, text, attrib, prevLeft)
                    if left > 200 and abs(pageWidth - (left + int(attrib['width'])) - int(attrib['left'])) < 50:
                            para_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            center = True
                    if left > pageWidth/2 and int(attrib['left']) + int(attrib['width']) > pageWidth - 300:
                        para_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    # else:
                    #     para_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    prevLeft = int(attrib['left']) + int(attrib['width'])
                    prevTop = int(attrib['top']) + int(attrib['height'])
            
            # if key == 'image':
            #     for item in value:
            #         attrib = item[1]
            #         img = attrib['src']
            #         mtch = re.match(r'_1.png',img) # avoid first image of complete page
            #         if(img):
            #             document.add_picture(img)            
        
        document.add_page_break()
    try:
        os.makedirs("Docx/" + fileName)
    except:
        pass
    document.save("Docx/" + fileName + "/" + fileName + ".docx")
    print("Completed: " + fileName)
